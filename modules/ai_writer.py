import os
import ollama
from docx import Document

def check_ollama_connection() -> bool:
    try:
        ollama.list()
        return True
    except:
        return False

def generate_doc(prompt: str, doc_name: str, save_path: str):
    if not check_ollama_connection():
        raise ConnectionError("Ollama no está conectado.")
    
    # Llamada a la API
    response = ollama.chat(
        model='llama3.2', 
        messages=[{'role': 'user', 'content': f"Escribe un texto detallado y académico sobre: {prompt}."}]
    )
    
    text_content = response['message']['content']
    
    # Nombre del archivo
    if not doc_name.endswith(".docx"):
        doc_name += ".docx"
    
    full_path = os.path.join(save_path, doc_name)
    
    # Creación del Word
    doc = Document()
    doc.add_heading(prompt.upper(), 0)
    doc.add_paragraph(text_content)
    doc.save(full_path)
    
    return full_path