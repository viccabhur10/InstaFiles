import customtkinter as ctk
from tkinter import filedialog, messagebox
import os
import shutil
import subprocess
import webbrowser 
from docx import Document
from PIL import Image
import ollama

# Configuración visual global
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

class InstaFiles(ctk.CTk):
    def __init__(self):
        super().__init__()

        # --- CONFIGURACIÓN DE VENTANA ---
        self.title("InstaFiles - Automatizador PC")
        self.geometry("950x750")
        
        try:
            self.iconbitmap(os.path.join("assets", "icono.ico"))
        except:
            pass

        # --- CARGA DE IMÁGENES ---
        self.img_mover_big = self.cargar_imagen("mover.png", alto_fijo=50)
        self.img_ia_big = self.cargar_imagen("ia.png", alto_fijo=50)
        self.img_foto_big = self.cargar_imagen("img.png", alto_fijo=50)
        
        # --- SEGURIDAD ---
        self.RUTAS_PROHIBIDAS = [
            "C:\\Windows", "C:\\Program Files", "C:\\Program Files (x86)", "/bin", "/usr", "/etc" 
        ]

        # --- PESTAÑAS ---
        self.tabview = ctk.CTkTabview(self, width=900, height=700)
        self.tabview.pack(padx=20, pady=20, fill="both", expand=True)

        self.tab_mover = self.tabview.add("Organizador")
        self.tab_ia = self.tabview.add("Redactor IA (Ollama)")
        self.tab_img = self.tabview.add("Conversor")

        self.configurar_modulo_mover()
        self.configurar_modulo_ia()
        self.configurar_modulo_imagenes()

    def cargar_imagen(self, nombre_archivo, alto_fijo=30):
      
        ruta = os.path.join("assets", nombre_archivo)
        if os.path.exists(ruta):
            img_original = Image.open(ruta)
            
            # Calculamos la proporción (Aspect Ratio)
            # Si la imagen es el doble de ancha que de alta, ratio será 2.
            ratio = img_original.width / img_original.height
            
            # El nuevo ancho será: Altura deseada * Proporción ( esto para que se vea bien y no estirada la imagen)
            nuevo_ancho = int(alto_fijo * ratio)
            
            return ctk.CTkImage(light_image=img_original, dark_image=img_original, size=(nuevo_ancho, alto_fijo))
        return None

    def es_ruta_segura(self, ruta):
        ruta_abs = os.path.abspath(ruta)
        for prohibida in self.RUTAS_PROHIBIDAS:
            if ruta_abs.lower().startswith(os.path.abspath(prohibida).lower()):
                return False
        return True

    def seleccionar_carpeta(self, entry_widget):
        ruta = filedialog.askdirectory()
        if ruta:
            entry_widget.delete(0, "end")
            entry_widget.insert(0, ruta)
            entry_widget.xview_moveto(1) 

    # =======================================================
    # MODULO 1: MOVER ARCHIVOS
    # =======================================================
    def configurar_modulo_mover(self):
        frame = self.tab_mover
        
        ctk.CTkLabel(frame, text=" Organización de Archivos", image=self.img_mover_big, compound="left", font=("Segoe UI", 26, "bold")).pack(pady=(20, 30))

        main_content = ctk.CTkFrame(frame, fg_color="transparent")
        main_content.pack(anchor="center")

        input_frame = ctk.CTkFrame(main_content)
        input_frame.pack(pady=10, padx=20, fill="x")
        input_frame.grid_columnconfigure(1, weight=1) 

        # Origen
        ctk.CTkLabel(input_frame, text="Origen:", font=("Arial", 14)).grid(row=0, column=0, padx=15, pady=15, sticky="e")
        self.entry_origen = ctk.CTkEntry(input_frame, width=400, height=35)
        self.entry_origen.grid(row=0, column=1, padx=10, sticky="we")
        ctk.CTkButton(input_frame, text="📂", width=50, height=35, font=("Segoe UI Emoji", 20), command=lambda: self.seleccionar_carpeta(self.entry_origen)).grid(row=0, column=2, padx=15)

        # Destino
        ctk.CTkLabel(input_frame, text="Destino:", font=("Arial", 14)).grid(row=1, column=0, padx=15, pady=15, sticky="e")
        self.entry_destino = ctk.CTkEntry(input_frame, width=400, height=35)
        self.entry_destino.grid(row=1, column=1, padx=10, sticky="we")
        ctk.CTkButton(input_frame, text="📂", width=50, height=35, font=("Segoe UI Emoji", 20), command=lambda: self.seleccionar_carpeta(self.entry_destino)).grid(row=1, column=2, padx=15)

        filter_frame = ctk.CTkFrame(main_content, fg_color="transparent")
        filter_frame.pack(pady=20, anchor="center")

        ctk.CTkLabel(filter_frame, text="Tipo de archivo (ej: .pdf):").grid(row=0, column=0, padx=10, sticky="e")
        self.entry_ext = ctk.CTkEntry(filter_frame, width=120)
        self.entry_ext.grid(row=0, column=1, padx=10)

        ctk.CTkLabel(filter_frame, text="Buscar por palabras (Opcional):").grid(row=0, column=2, padx=10, sticky="e")
        self.entry_key = ctk.CTkEntry(filter_frame, width=250)
        self.entry_key.grid(row=0, column=3, padx=10)

        ctk.CTkButton(frame, text="MOVER ARCHIVOS", fg_color="#e63946", hover_color="#d62828", height=50, font=("Arial", 16, "bold"), command=self.accion_mover).pack(pady=30)

    def accion_mover(self):
        origen = self.entry_origen.get()
        destino = self.entry_destino.get()
        ext = self.entry_ext.get()
        keyword = self.entry_key.get()

        if not self.es_ruta_segura(origen) or not self.es_ruta_segura(destino):
            messagebox.showerror("Seguridad", "Ruta protegida detectada.")
            return
        if not os.path.exists(origen): return
        
        archivos_a_mover = []
        try:
            for f in os.listdir(origen):
                if not ext or f.lower().endswith(ext.lower()):
                    if keyword and keyword.lower() not in f.lower(): continue
                    archivos_a_mover.append(f)
        except Exception as e: 
            messagebox.showerror("Error", f"Error carpeta: {e}")
            return

        if not archivos_a_mover:
            messagebox.showinfo("Info", "No se encontraron archivos.")
            return

        if messagebox.askyesno("Confirmar", f"¿Mover {len(archivos_a_mover)} archivos a:\n{destino}?"):
            if not os.path.exists(destino): os.makedirs(destino)
            movidos = 0
            for f in archivos_a_mover:
                try: 
                    shutil.move(os.path.join(origen, f), os.path.join(destino, f))
                    movidos += 1
                except: pass
            messagebox.showinfo("Hecho", f"Se movieron {movidos} archivos.")

    # =======================================================
    # MODULO 2: IA LOCAL
    # =======================================================
    def configurar_modulo_ia(self):
        frame = self.tab_ia
        
        header_frame = ctk.CTkFrame(frame, fg_color="transparent")
        header_frame.pack(pady=(20, 10), fill="x", padx=20)
        ctk.CTkLabel(header_frame, text=" Redactor IA Local", image=self.img_ia_big, compound="left", font=("Segoe UI", 26, "bold")).pack(side="left")
        self.lbl_estado_ollama = ctk.CTkLabel(header_frame, text="Comprobando...", font=("Arial", 14, "bold"))
        self.lbl_estado_ollama.pack(side="right", padx=10)
        
        self.config_frame = ctk.CTkFrame(frame, border_width=1, border_color="#E07A5F")
        self.config_frame.pack(pady=10, padx=40, fill="x")
        
        ctk.CTkLabel(self.config_frame, text="Configuración Manual (NO TOCAR si arriba dice CONECTADO):", font=("Arial", 11, "bold"), text_color="#E07A5F").pack(anchor="w", padx=15, pady=(10, 5))
        
        path_frame = ctk.CTkFrame(self.config_frame, fg_color="transparent")
        path_frame.pack(fill="x", padx=10, pady=(0, 10))
        self.entry_ollama_path = ctk.CTkEntry(path_frame, placeholder_text="Ruta al .exe de Ollama")
        self.entry_ollama_path.pack(side="left", fill="x", expand=True, padx=5)
        ctk.CTkButton(path_frame, text="Buscar .exe", width=90, command=self.buscar_exe_ollama).pack(side="left", padx=5)
        ctk.CTkButton(path_frame, text="Forzar Inicio", width=100, fg_color="#E07A5F", hover_color="#D06A4F", command=self.iniciar_ollama_manual).pack(side="left", padx=5)
        
        btn_help = ctk.CTkButton(frame, text="¿No tienes Ollama instalado? Descárgalo aquí", fg_color="transparent", text_color="#4cc9f0", hover=False, command=lambda: webbrowser.open("https://ollama.com/download"))
        btn_help.pack(pady=(0, 20))

        work_frame = ctk.CTkFrame(frame)
        work_frame.pack(padx=40, fill="x", pady=10)
        work_frame.grid_columnconfigure(1, weight=1)

        ctk.CTkLabel(work_frame, text="Nombre archivo:", font=("Arial", 14)).grid(row=0, column=0, padx=15, pady=15, sticky="e")
        self.entry_doc_name = ctk.CTkEntry(work_frame, height=35, placeholder_text="Ej: Trabajo_Historia")
        self.entry_doc_name.grid(row=0, column=1, padx=10, sticky="we", columnspan=2)

        # Ruta
        ctk.CTkLabel(work_frame, text="Guardar en:", font=("Arial", 14)).grid(row=1, column=0, padx=15, pady=15, sticky="e")
        self.entry_doc_path = ctk.CTkEntry(work_frame, height=35, placeholder_text="Selecciona una carpeta...")
        self.entry_doc_path.grid(row=1, column=1, padx=10, sticky="we")
        ctk.CTkButton(work_frame, text="📂", width=50, height=35, font=("Segoe UI Emoji", 20), command=lambda: self.seleccionar_carpeta(self.entry_doc_path)).grid(row=1, column=2, padx=15)

        ctk.CTkLabel(frame, text="Prompt (Instrucciones para la IA):", font=("Arial", 14)).pack(anchor="w", padx=40, pady=(15, 5))
        self.text_prompt = ctk.CTkTextbox(frame, height=120)
        self.text_prompt.pack(fill="x", padx=40, pady=5)

        self.btn_generar = ctk.CTkButton(frame, text="GENERAR DOCUMENTO", fg_color="#2a9d8f", hover_color="#21867a", height=50, font=("Arial", 16, "bold"), command=self.accion_generar_ia)
        self.btn_generar.pack(pady=20)

        self.verificar_ollama()

    def buscar_exe_ollama(self):
        archivo = filedialog.askopenfilename(filetypes=[("Ejecutables", "*.exe")])
        if archivo:
            self.entry_ollama_path.delete(0, "end")
            self.entry_ollama_path.insert(0, archivo)

    def verificar_ollama(self):
        try:
            ollama.list()
            self.lbl_estado_ollama.configure(text="🟢 CONECTADO", text_color="#2dc653")
            self.ollama_ready = True
        except:
            self.lbl_estado_ollama.configure(text="🔴 DESCONECTADO", text_color="#d00000")
            self.ollama_ready = False

    def iniciar_ollama_manual(self):
        ruta_exe = self.entry_ollama_path.get()
        if os.path.exists(ruta_exe) and ruta_exe.endswith("ollama.exe"):
            try:
                subprocess.Popen([ruta_exe, "serve"], creationflags=subprocess.CREATE_NO_WINDOW)
                messagebox.showinfo("Iniciando", "Espera 5-10 segundos y el estado debería cambiar.")
                self.after(8000, self.verificar_ollama) 
            except Exception as e: messagebox.showerror("Error", f"Error: {e}")
        else: messagebox.showwarning("Ruta inválida", "Selecciona 'ollama.exe'.")

    def accion_generar_ia(self):
        self.verificar_ollama()
        if not self.ollama_ready:
            messagebox.showerror("Error", "Ollama no está conectado.")
            return
        prompt = self.text_prompt.get("1.0", "end-1c")
        nombre = self.entry_doc_name.get()
        ruta_guardado = self.entry_doc_path.get()
        if not nombre or len(prompt) < 5 or not ruta_guardado:
            messagebox.showwarning("Faltan datos", "Rellena todos los campos.")
            return
        if not os.path.exists(ruta_guardado):
            messagebox.showerror("Error", "La ruta no existe.")
            return
        self.btn_generar.configure(state="disabled", text="Pensando...")
        self.update()
        try:
            response = ollama.chat(model='llama3.2', messages=[{'role': 'user', 'content': f"Escribe un texto detallado sobre: {prompt}. Formato académico."}])
            texto = response['message']['content']
            if not nombre.endswith(".docx"): nombre += ".docx"
            full_path = os.path.join(ruta_guardado, nombre)
            doc = Document()
            doc.add_heading(prompt.upper(), 0)
            doc.add_paragraph(texto)
            doc.save(full_path)
            messagebox.showinfo("Éxito", f"Guardado en:\n{full_path}")
        except Exception as e: messagebox.showerror("Error", f"Error: {e}")
        finally: self.btn_generar.configure(state="normal", text="GENERAR DOCUMENTO")


    # =======================================================
    # MODULO 3: IMÁGENES
    # =======================================================
    def configurar_modulo_imagenes(self):
        frame = self.tab_img
        
        ctk.CTkLabel(frame, text=" Conversor de Formatos", image=self.img_foto_big, compound="left", font=("Segoe UI", 26, "bold")).pack(pady=(20, 40))

        main_content = ctk.CTkFrame(frame, fg_color="transparent")
        main_content.pack(anchor="center", fill="x", padx=50)
        main_content.grid_columnconfigure(1, weight=1)

        ctk.CTkLabel(main_content, text="Seleccionar imagen:", font=("Arial", 14)).grid(row=0, column=0, padx=15, pady=15, sticky="e")
        self.entry_img_path = ctk.CTkEntry(main_content, height=35, placeholder_text="Ninguna imagen seleccionada", state="readonly")
        self.entry_img_path.grid(row=0, column=1, padx=10, sticky="we")
        
        # Botón
        ctk.CTkButton(main_content, text="📂", width=50, height=35, font=("Segoe UI Emoji", 20), command=self.seleccionar_imagen).grid(row=0, column=2, padx=15)

        format_frame = ctk.CTkFrame(frame, fg_color="transparent")
        format_frame.pack(pady=20)
        ctk.CTkLabel(format_frame, text="Convertir a formato:", font=("Arial", 14)).pack(side="left", padx=15)
        self.combo_formato = ctk.CTkComboBox(format_frame, values=["PDF", "PNG", "JPEG", "WEBP", "ICO"], width=150, height=35, font=("Arial", 14))
        self.combo_formato.set("PDF")
        self.combo_formato.pack(side="left")
        
        ctk.CTkButton(frame, text="CONVERTIR IMAGEN", fg_color="#0077b6", hover_color="#0096c7", height=50, font=("Arial", 16, "bold"), command=self.accion_convertir_img).pack(pady=30)
        self.img_seleccionada_full_path = None

    def seleccionar_imagen(self):
        archivo = filedialog.askopenfilename(filetypes=[("Imágenes", "*.jpg *.png *.jpeg *.webp *.bmp")])
        if archivo:
            self.img_seleccionada_full_path = archivo
            self.entry_img_path.configure(state="normal")
            self.entry_img_path.delete(0, "end")
            self.entry_img_path.insert(0, archivo)
            self.entry_img_path.xview_moveto(1)
            self.entry_img_path.configure(state="readonly")
            
    def accion_convertir_img(self):
        if not self.img_seleccionada_full_path or not os.path.exists(self.img_seleccionada_full_path):
            messagebox.showwarning("Atención", "Selecciona una imagen primero.")
            return
        formato = self.combo_formato.get()
        ruta_guardado = filedialog.asksaveasfilename(defaultextension=f".{formato.lower()}", filetypes=[(formato, f"*.{formato.lower()}")], title="Guardar como...")
        if ruta_guardado:
            try:
                img = Image.open(self.img_seleccionada_full_path)
                if img.mode in ("RGBA", "P") and formato in ["JPEG", "PDF"]:
                     img = img.convert('RGB')
                if formato == "ICO":
                    img.save(ruta_guardado, format='ICO', sizes=[(256, 256)])
                else:
                    img.save(ruta_guardado)
                messagebox.showinfo("Éxito", f"Imagen convertida a {formato}.")
            except Exception as e: messagebox.showerror("Error", f"Error: {str(e)}")

if __name__ == "__main__":
    app = InstaFiles()
    app.mainloop()